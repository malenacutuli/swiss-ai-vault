"""
DOCX document generator using python-docx.

Generates formatted Word documents with:
- Title and heading styles
- Sections with content
- Tables with styling
- Paragraphs with formatting
"""

import os
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .base import DocumentGenerator, DocumentContent


class DOCXGenerator(DocumentGenerator):
    """Generator for Microsoft Word (.docx) documents."""

    def __init__(self, content: DocumentContent):
        """
        Initialize DOCX generator.

        Args:
            content: DocumentContent with title and sections
        """
        super().__init__(content)
        self.document = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles."""
        # Title style
        styles = self.document.styles
        if 'CustomTitle' not in [s.name for s in styles]:
            title_style = styles.add_style('CustomTitle', 1)  # 1 = PARAGRAPH
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0, 0, 139)

    def get_format_name(self) -> str:
        """Return format name."""
        return "docx"

    def validate_content(self) -> bool:
        """
        Validate content structure for DOCX generation.

        Returns:
            bool: True if valid

        Raises:
            ValueError: If content is invalid
        """
        if not self.content.title:
            raise ValueError("Document title is required")

        if not self.content.sections:
            raise ValueError("At least one section is required")

        for idx, section in enumerate(self.content.sections):
            if "heading" not in section:
                raise ValueError(f"Section {idx} missing 'heading' field")

        return True

    def generate(self, output_path: str) -> str:
        """
        Generate DOCX document.

        Args:
            output_path: Path to save the document

        Returns:
            str: Path to generated document
        """
        self.validate_content()

        # Add title
        title = self.document.add_heading(self.content.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata if present
        if self.content.metadata:
            author = self.content.metadata.get("author", "Swiss AI Vault Agent")
            self.document.core_properties.author = author
            if "created_at" in self.content.metadata:
                created_at = self.content.metadata["created_at"]
                meta_para = self.document.add_paragraph()
                meta_para.add_run(f"Created: {created_at}").italic = True
                meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.document.add_paragraph()  # Spacing

        # Add sections
        for section in self.content.sections:
            self._add_section(section)

        # Save document
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.document.save(output_path)

        # Update metadata
        file_size = os.path.getsize(output_path)
        self.update_metadata(file_size=file_size)

        return output_path

    def _add_section(self, section: Dict[str, Any]):
        """
        Add a section to the document.

        Args:
            section: Section data with heading and content
        """
        # Add section heading
        heading = section.get("heading", "Section")
        self.document.add_heading(heading, level=1)

        # Add content based on type
        content_type = section.get("type", "text")

        if content_type == "text":
            content = section.get("content", "")
            if isinstance(content, list):
                for item in content:
                    self.document.add_paragraph(str(item))
            else:
                self.document.add_paragraph(str(content))

        elif content_type == "bullet_list":
            items = section.get("items", [])
            for item in items:
                self.document.add_paragraph(str(item), style='List Bullet')

        elif content_type == "numbered_list":
            items = section.get("items", [])
            for item in items:
                self.document.add_paragraph(str(item), style='List Number')

        elif content_type == "table":
            self._add_table(section)

        elif content_type == "code":
            code_content = section.get("content", "")
            code_para = self.document.add_paragraph(code_content)
            code_para.style = 'Normal'
            run = code_para.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

        # Add spacing after section
        self.document.add_paragraph()

    def _add_table(self, section: Dict[str, Any]):
        """
        Add a table to the document.

        Args:
            section: Section data with table configuration
        """
        headers = section.get("headers", [])
        rows = section.get("rows", [])

        if not headers or not rows:
            return

        # Create table
        table = self.document.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'

        # Add headers
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            header_cells[idx].text = str(header)
            # Make header bold
            for paragraph in header_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Add data rows
        for row_idx, row_data in enumerate(rows, start=1):
            row_cells = table.rows[row_idx].cells
            for col_idx, cell_value in enumerate(row_data):
                if col_idx < len(row_cells):
                    row_cells[col_idx].text = str(cell_value)
